import os
from datetime import datetime

import google.generativeai as genai
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.security import HTTPBearer
from firebase import get_auth, get_db
from pydantic import BaseModel

# while starting the application, it loads the .env files
load_dotenv()

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# creates the fastapi app, given the title as well
app = FastAPI(title="AI Resume Analyser", version="1.0.0")

# register an authentication scheme globally, later we use to protect the endpoint
security = HTTPBearer()


# using basemodel , we can return the structured output,
# it is not allowing the user to write the wrong queries
class Analysis(BaseModel):
    score: float
    skills: list
    gaps: list
    suggestions: list


@app.post("/register")
async def register(email: str, password: str):
    try:
        auth_client = get_auth()
        user = auth_client.create_user(email=email, password=password)
        return {"user_id": user.uid, "email": user.email, "message": "User created"}
    except Exception as e:
        return {"error": str(e)}


@app.post("/login")
async def login(email: str, password: str):
    try:
        auth_client = get_auth()
        user = auth_client.get_user_by_email(email)

        custom_token = auth_client.create_custom_token(user.uid)
        return {
            "token": custom_token.decode(),
            "user id": user.uid,
            "email": user.email,
        }
    except Exception:
        return {"error": "Invalid credentials"}


# async def verify_firebase_token(credentials=Depends(security)):
#     token = credentials.credentials
#     print(f"🔍 TOKEN LENGTH: {len(token)}")
#     print(f"🔍 TOKEN START: {token[:20]}...")
#     print(f"🔍 TOKEN END: ...{token[-20:]}")

#     # TEMP: Accept ANY token for 5 mins
#     print("🔓 DEBUG: Bypassing token check")
#     return {"uid": "debug-wk2C98jDMuWpvbIKQEnE7BCxNfQ2"}


@app.get("/")
async def root():
    return {"message": "AI Resume Analyser", "status": "live"}


@app.get("/health")
async def health():
    return {"status": "healthy", "gcp_ready": True}


# using File - > accepts Files/Forms
# Depens works for dependency injection
# token -> extract authentication bearer and inject token object automatically


# @app.post("/analyse-resume")
# async def analyse_resume(
#     file: UploadFile = File(...), user=Depends(verify_firebase_token)
# ):
#     # FILE TYPE CHECK
#     allowed_types = [".pdf", ".docx"]
#     if not any(file.filename.lower().endswith(ext) for ext in allowed_types):
#         raise HTTPException(400, "Only PDF and DOCX files supported")

#     content = ""

#     try:
#         # PDF HANDLING
#         if file.filename.lower().endswith(".pdf"):
#             pdf_reader = PyPDF2.PdfReader(file.file)
#             for page in pdf_reader.pages:
#                 content += page.extract_text()

#         # DOCX HANDLING WITH ERROR CATCHING
#         elif file.filename.lower().endswith(".docx"):
#             try:
#                 # Reset file pointer for DOCX
#                 file.file.seek(0)
#                 doc = Document(file.file)
#                 for para in doc.paragraphs:
#                     content += para.text + "\n"
#             except Exception as docx_error:
#                 print(f"DOCX Error: {docx_error}")
#                 raise HTTPException(400, f"Invalid DOCX file: {str(docx_error)[:100]}")

#     except Exception as parse_error:
#         raise HTTPException(400, f"File parsing failed: {str(parse_error)[:100]}")

#     # # SCORING (your fixed base:30)
#     # skills = ["FastAPI", "Docker", "Python", "AWS", "Kubernetes", "GCP", "Linux"]
#     # found_skills = [s for s in skills if s.lower() in content.lower()]
#     # word_count = len(content.split())

#     # base_score = 30
#     # skills_bonus = len(found_skills) * 8
#     # length_score = min(word_count / 10, 20)

#     # analysis = {
#     #     "score": round(min(base_score + skills_bonus + length_score, 100), 2),
#     #     "skills_found": found_skills,
#     #     "word_count": word_count,
#     #     "filename": file.filename,
#     #     "timestamp": datetime.now().isoformat(),
#     # }

#     # # FIREBASE SAVE (unchanged)
#     uid = user["uid"]
#     db = get_db()
#     user_ref = db.child(f"users/{uid}/analyses")
#     analysis_id = user_ref.push(analysis)

#     # return {
#     #     "analysis_id": analysis_id.key,
#     #     "data": analysis,
#     #     "firebase_path": f"users/{uid}/analyses/{analysis_id.key}",
#     # }

#     # REAL AI ANALYSIS
#     model = genai.GenerativeModel("gemini-1.5-flash")
#     prompt = f"""
#     Analyze this resume for SDE-2 Backend roles:
#     {content}

#     Return VALID JSON only:
#     {{
#         "score": 85,
#         "skills": ["FastAPI", "Docker", "GCP"],
#         "strengths": ["Production deployment experience"],
#         "gaps": ["Kubernetes", "Terraform"],
#         "suggestions": ["Add GCP projects to portfolio"]
#     }}
#     """

#     response = model.generate_content(prompt)
#     analysis = json.loads(response.text)

#     return {
#         "analysis_id": analysis_id.key,
#         "data": analysis,
#         "firebase_path": f"users/{uid}/analyses/{analysis_id.key}",
#     }


# @app.post("/analyse-resume")
# async def analyse_resume(
#     file: UploadFile = File(...), user=Depends(verify_firebase_token)
# ):
@app.post("/analyse-resume")
async def analyse_resume(file: UploadFile = File(...)):
    import json
    import os

    # VALIDATE FILE TYPE
    if not file.filename.lower().endswith((".pdf", ".docx")):
        raise HTTPException(400, "PDF/DOCX only")

    # PARSE FILE CONTENT
    content = ""
    try:
        if file.filename.lower().endswith(".pdf"):
            import fitz  # pip install pymupdf

            pdf_bytes = await file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page in doc:
                content += page.get_text()
            doc.close()
        else:  # docx
            from docx import Document

            doc = Document(file.file)
            content += "\n".join(p.text for p in doc.paragraphs)

        if len(content.strip()) < 50:
            raise ValueError("Extracted text too short - parsing failed")

    except Exception as e:
        raise HTTPException(400, f"Parse failed: {str(e)}")

    # GEMINI AI ANALYSIS - BULLETPROOF
    try:
        import google.generativeai as genai

        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key or len(api_key) < 20:
            raise ValueError("Missing GEMINI_API_KEY")

        # CREATE PRECISE PROMPT
        prompt = f"""
        Analyze this resume for SDE-2 Backend Engineer role (FastAPI, Docker, GCP, Kubernetes, Python):
        {content[:4000]}
        
        Return ONLY valid JSON with this EXACT structure - no explanations:
        {{
            "score": 85.5,
            "skills": ["FastAPI", "Docker", "Python"],
            "strengths": ["3+ years production experience"],
            "gaps": ["Kubernetes missing", "No Terraform"],
            "suggestions": ["Add GCP Cloud Run deployment experience", "Get Kubernetes certification"]
        }}
        """

        genai.configure(api_key=api_key)

        # DEBUG: LIST AVAILABLE MODELS (FIRST TIME ONLY)
        print("🔍 Available Gemini models:")
        for model in genai.list_models():
            if "generateContent" in model.supported_generation_methods:
                print(f"  - {model.name}")

        # TRY MODELS IN ORDER (YOUR ORIGINAL FIRST)
        working_models = [
            "gemini-2.5-flash",  # ✅ TOP OF YOUR LIST - STABLE
            "gemini-2.5-pro",  # ✅ SECOND - STABLE
            "gemini-flash-latest",  # ✅ ALWAYS WORKS
            "gemini-pro",  # ✅ FINAL FALLBACK
        ]

        analysis = None
        for model_name in working_models:
            try:
                print(f"🔄 Trying model: {model_name}")
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(prompt)
                print(f"🌟 RAW RESPONSE: {response.text[:150]}...")

                # STRICT JSON VALIDATION
                cleaned_response = response.text.strip()
                if cleaned_response.startswith("```json"):
                    cleaned_response = cleaned_response[7:-3].strip()
                elif cleaned_response.startswith("```"):
                    cleaned_response = cleaned_response[3:].strip()

                analysis = json.loads(cleaned_response)
                print(f"✅ SUCCESS with {model_name}")
                break

            except Exception as model_error:
                print(f"❌ {model_name} failed: {str(model_error)[:100]}")
                continue

        if not analysis:
            raise Exception("No working Gemini models available")

    except Exception as ai_error:
        print(f"💥 GEMINI FAILED: {ai_error}")
        raise HTTPException(500, f"AI analysis failed: {str(ai_error)}")

    # ENRICH ANALYSIS WITH METRICS
    analysis.update(
        {
            "word_count": len(content.split()),
            "filename": file.filename,
            "timestamp": datetime.now().isoformat(),
            "raw_content_length": len(content),
        }
    )

    # SAVE TO FIRESTORE (YOUR WORKING CODE)
    try:
        uid = "demo-user-123"  # ← DUMMY UID - NO AUTH NEEDED
        db_client = get_db()
        user_ref = db_client.child(f"users/{uid}/analyses")
        analysis_id = user_ref.push(analysis)

        return {
            "success": True,
            "analysis_id": analysis_id.key,
            "data": analysis,
            "firebase_path": f"users/{uid}/analyses/{analysis_id.key}",
        }
    except Exception:
        return {"success": False, "data": analysis}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
